import os
import logging
from typing import List
import docx
from core.pdf_reader import ExtractedElement, PDFReader

logger = logging.getLogger("rag-service.docx-reader")

class DocxReader:
    """
    Docx Reader designed to ingest Vietnamese Literature textbook documents in .docx format.
    Reuses PDFReader's heuristic detection for headings and items.
    """
    def __init__(self):
        # We instantiate a PDFReader just to reuse its heuristic methods 
        # (_is_heading, _is_numbered_item, _clean_text) easily without copying code.
        self._pdf_reader = PDFReader()

    def read(self, file_path: str) -> List[ExtractedElement]:
        if not os.path.exists(file_path):
            raise FileNotFoundError(f"Docx file not found at: {file_path}")

        source_file = os.path.basename(file_path)
        logger.info(f"Starting extraction for Docx file: {source_file}")

        elements: List[ExtractedElement] = []
        doc = docx.Document(file_path)

        for paragraph in doc.paragraphs:
            text = paragraph.text.strip()
            if not text:
                continue

            cleaned_text = self._pdf_reader._clean_text(text, is_final=False)
            if not cleaned_text:
                continue

            style_name = paragraph.style.name.lower()
            is_list_style = 'list' in style_name or 'bullet' in style_name
            
            # Detect Word's internal numbering/bullet formatting via XML.
            # This catches bullets/lists even when the style name is "Normal".
            has_word_numbering = False
            pPr = paragraph._element.find(
                '{http://schemas.openxmlformats.org/wordprocessingml/2006/main}pPr'
            )
            if pPr is not None:
                numPr = pPr.find(
                    '{http://schemas.openxmlformats.org/wordprocessingml/2006/main}numPr'
                )
                if numPr is not None:
                    has_word_numbering = True
            
            # Determine type. Order matters!
            # 1. Word Heading Style
            # 2. Numbered Item (1. Tác giả)
            # 3. List Item (• Phần 1)
            # 4. Heading Heuristic (Tác giả, or Phần 1 if short enough)
            # 5. Word numbering fallback
            if 'heading' in style_name:
                el_type = "heading"
            elif self._pdf_reader._is_numbered_item(cleaned_text):
                el_type = "numbered_item"
            elif self._pdf_reader._is_list_item(cleaned_text):
                el_type = "list"
            elif is_list_style or has_word_numbering:
                el_type = "list"
            elif self._pdf_reader._is_heading(cleaned_text):
                el_type = "heading"
            else:
                el_type = "paragraph"

            # Re-clean as final for the actual element text
            final_text = self._pdf_reader._clean_text(cleaned_text, is_final=True)
            
            elements.append(
                ExtractedElement(
                    page=1, # Docx doesn't map to pages easily here
                    type=el_type,
                    raw_text=final_text,
                    source_file=source_file
                )
            )

        # Extract Tables
        for table in doc.tables:
            formatted_rows = []
            for row in table.rows:
                row_data = [self._pdf_reader._clean_text(cell.text.strip()) for cell in row.cells]
                if any(row_data):
                    formatted_rows.append(" | ".join(row_data))
            
            if formatted_rows:
                table_text = "\n".join(formatted_rows)
                elements.append(
                    ExtractedElement(
                        page=1,
                        type="table",
                        raw_text=table_text,
                        source_file=source_file
                    )
                )

        logger.info(f"Successfully extracted {len(elements)} elements using DocxReader.")
        return elements
